import logging
import tempfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from rdkit import Chem
    from rdkit.Chem import Draw
    _HAS_RDKIT = True
except ImportError:
    _HAS_RDKIT = False

logger = logging.getLogger(__name__)


class ReportService:
    def generate(
        self,
        request_id: str,
        predictions: dict,
        interpretation: dict,
        smiles_list: list[str],
    ) -> str:
        """Generate a Word (.docx) ADMET report and return the file path."""
        doc = Document()

        # -- Styles --
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        # -- Title --
        title = doc.add_heading("Analysis Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

        # -- Metadata --
        doc.add_paragraph("")
        meta_table = doc.add_table(rows=3, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_cells = [
            ("Report ID", request_id),
            ("Generated", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")),
            ("Molecules", str(len(smiles_list))),
        ]
        for i, (label, value) in enumerate(meta_cells):
            meta_table.rows[i].cells[0].text = label
            meta_table.rows[i].cells[1].text = value
            for cell in meta_table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    paragraph.style.font.size = Pt(10)

        doc.add_paragraph("")

        # -- AI Overview --
        ai_overview = interpretation.get("ai_overview", "")
        if ai_overview:
            doc.add_heading("AI Overview", level=1)
            p = doc.add_paragraph()
            run = p.add_run(ai_overview)
            run.font.size = Pt(11)
            doc.add_paragraph("")

        # -- Per-molecule sections --
        for smiles in smiles_list:
            doc.add_heading(f"Molecule: {smiles}", level=1)

            # 2D structure image (requires rdkit)
            if _HAS_RDKIT:
                mol = Chem.MolFromSmiles(smiles)
                if mol:
                    try:
                        img = Draw.MolToImage(mol, size=(400, 300))
                        buf = BytesIO()
                        img.save(buf, format="PNG")
                        buf.seek(0)
                        paragraph = doc.add_paragraph()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run()
                        run.add_picture(buf, width=Inches(3.5))
                    except Exception as e:
                        logger.warning("Could not render molecule image for %s: %s", smiles, e)
                        doc.add_paragraph(f"[Could not render 2D structure for {smiles}]")
            else:
                doc.add_paragraph(f"[2D structure rendering unavailable — install rdkit]")

            # ADMET results table
            mol_preds = predictions.get(smiles, {})
            if mol_preds:
                doc.add_heading("ADMET Predictions", level=2)

                # Group results by category
                groups = {}
                for endpoint_name, data in mol_preds.items():
                    group = data.get("group", "other")
                    if group not in groups:
                        groups[group] = []
                    groups[group].append((endpoint_name, data))

                for group_name in ["absorption", "distribution", "metabolism", "excretion", "toxicity"]:
                    if group_name not in groups:
                        continue

                    doc.add_heading(group_name.title(), level=3)

                    table = doc.add_table(rows=1, cols=5)
                    table.style = "Light Grid Accent 1"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    headers = ["Endpoint", "Value", "Unit", "Std Dev", "Type"]
                    for i, h in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = h
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    for endpoint_name, data in groups[group_name]:
                        row = table.add_row().cells
                        row[0].text = data.get("display_name", endpoint_name)
                        row[1].text = f"{data.get('value', 'N/A')}"
                        row[2].text = data.get("unit", "") or ""
                        row[3].text = f"\u00b1 {data.get('std', 'N/A')}"
                        row[4].text = data.get("type", "")

        # -- Confidence & Uncertainty --
        doc.add_heading("Model Confidence", level=1)
        doc.add_paragraph(
            "Predictions are generated from a 5-model ensemble. The standard deviation "
            "(Std Dev) column indicates prediction uncertainty across ensemble members. "
            "Lower values indicate higher model confidence in the prediction."
        )

        # -- AI Interpretation --
        doc.add_heading("AI Interpretation", level=1)

        executive = interpretation.get("executive_summary", "")
        if executive:
            p = doc.add_paragraph()
            run = p.add_run(executive)
            run.font.size = Pt(11)

        high_risk = interpretation.get("high_risk_flags", [])
        if high_risk:
            doc.add_heading("High Risk Flags", level=2)
            for flag in high_risk:
                p = doc.add_paragraph(flag, style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        admet_interp = interpretation.get("admet_interpretation", [])
        if admet_interp:
            doc.add_heading("Detailed Endpoint Analysis", level=2)
            for item in admet_interp:
                endpoint = item.get("endpoint", "Unknown")
                value = item.get("value", "")
                assessment = item.get("assessment", "")
                explanation = item.get("explanation", "")

                p = doc.add_paragraph()
                run = p.add_run(f"{endpoint}: ")
                run.bold = True
                p.add_run(f"{value} ")

                # Color-code assessment
                assessment_run = p.add_run(f"[{assessment}]")
                if assessment == "concerning":
                    assessment_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                elif assessment == "moderate":
                    assessment_run.font.color.rgb = RGBColor(0xEA, 0xB3, 0x08)
                else:
                    assessment_run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

                if explanation:
                    p.add_run(f" — {explanation}")

        next_steps = interpretation.get("recommended_next_steps", [])
        if next_steps:
            doc.add_heading("Recommended Next Steps", level=2)
            for step in next_steps:
                doc.add_paragraph(step, style="List Bullet")

        # -- Disclaimer --
        doc.add_heading("Disclaimer", level=1)
        disclaimer = interpretation.get(
            "disclaimer",
            "These are computational predictions and must be validated experimentally "
            "before making any drug development decisions.",
        )
        p = doc.add_paragraph(disclaimer)
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)

        # -- Save --
        report_dir = Path(tempfile.mkdtemp())
        report_path = report_dir / f"analysis_report_{request_id[:8]}.docx"
        doc.save(str(report_path))

        logger.info("Report generated: %s", report_path)
        return str(report_path)
